from sympy import *
from docx import Document
from docx.shared import Inches
import os
from lxml import etree
import abc


x = [
    [
        ('$x^{2} - 6 x + 8$', ['Factored Form: ']),
        ('$- x^{2} + 6 x - 8$', ['Factored Form: ']),
        ('$x^{2} - 49$', ['Factored Form: ']),
        ('$- x^{2} - 6 x + 27$', ['Factored Form: ']),
        ('$- x^{2} - 7 x + 8$', ['Factored Form: ']),
        ('$- 3 x^{2} + 23 x - 14$', ['Factored Form: ']),
        ('$2 x^{2} - 3 x - 14$', ['Factored Form: ']),
        ('$- 4 x^{2} - 12 x + 7$', ['Factored Form: ']),
        ('$2 x^{2} - 3 x - 35$', ['Factored Form: '])
    ],

    [('$x^{2} - 6 x + 8$', ['Factored Form: $\\left(2 - x\\right) \\left(4 - x\\right)$']),
     ('$- x^{2} + 6 x - 8$',
      ['Factored Form: $\\left(4 - x\\right) \\left(x - 2\\right)$']),
     ('$x^{2} - 49$',
      ['Factored Form: $\\left(x - 7\\right) \\left(x + 7\\right)$']),
     ('$- x^{2} - 6 x + 27$',
      ['Factored Form: $\\left(3 - x\\right) \\left(x + 9\\right)$']),
     ('$- x^{2} - 7 x + 8$',
      ['Factored Form: $\\left(1 - x\\right) \\left(x + 8\\right)$']),
     ('$- 3 x^{2} + 23 x - 14$',
      ['Factored Form: $\\left(7 - x\\right) \\left(3 x - 2\\right)$']),
     ('$2 x^{2} - 3 x - 14$',
      ['Factored Form: $\\left(x + 2\\right) \\left(2 x - 7\\right)$']),
     ('$- 4 x^{2} - 12 x + 7$',
      ['Factored Form: $\\left(1 - 2 x\\right) \\left(2 x + 7\\right)$']),
     ('$2 x^{2} - 3 x - 35$',
      ['Factored Form: $\\left(5 - x\\right) \\left(- 2 x - 7\\right)$'])
     ]
]

# def mathMlConvertion(expr1):
#     expr1xml = mathml(expr1, printer='presentation')
#     tree = etree.fromstring(
#         '<math xmlns="http://www.w3.org/1998/Math/MathML">'+expr1xml+'</math>')

#     # convert to MS Office structure
#     xslt = etree.parse('MML2OMML.XSL')
#     transform = etree.XSLT(xslt)
#     new_dom = transform(tree)

#     return new_dom


# def worksheet(details, filename):

#     for i in range(2):

#         filename = filename+'anskey' if i == 1 else filename
#         document = Document()

#         document.add_heading('Worksheet 1', 0)
#         for each in details:
#             mainTopic, subTopic, instruction, items, givQuesAns = each
#             document.add_paragraph(subTopic)
#             document.add_paragraph(instruction)
#             for item in givQuesAns[i]:
#                 print(item[0])
#                 expr1 = mathMlConvertion(
#                     item[0]) if item[0] != str else item[0]
#                 document.add_paragraph(item[0], style='List Bullet')

#     # p.add_run('bold').bold = True
#     # p.add_run(' and some ')
#     # p.add_run('italic.').italic = True

#     # document.add_picture('monty-truth.png', width=Inches(1.25))

#     document.add_page_break()

#     document.save(filename)


# def main(details):
#     # print(details)
#     # checks if the same filename exists in the directory
#     files = os.listdir('output')

#     i = 1
#     temp = 'pol'
#     while temp+'.docx' in files:
#         temp = 'pol'
#         temp += str(i)
#         i += 1
#     filename = 'output/' + temp
#     worksheet(details, filename)

document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

# document.add_picture('monty-truth.png', width=Inches(1.25))

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = x[0][0]
    row_cells[1].text = x[0][1]
    row_cells[2].text = x[0][2]

document.add_page_break()

document.save('demo.docx')
